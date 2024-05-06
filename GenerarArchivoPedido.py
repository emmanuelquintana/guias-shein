import os
import json
from datetime import date
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION

def select_folder():
    """Permite al usuario seleccionar una carpeta donde se encuentran las imágenes."""
    import tkinter as tk
    from tkinter import filedialog

    root = tk.Tk()
    root.withdraw()
    folder_selected = filedialog.askdirectory()
    return folder_selected

def create_orders_doc(data, base_image_folder):
    """Crea un archivo Word con la información de los productos y las imágenes correspondientes."""
    doc = Document()

    # Cambiar orientación del documento a horizontal
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE

    # Agregar título y subtítulo
    today = date.today().strftime("%d/%m/%Y")
    doc.add_heading(f"Pedidos Shein - {today}", level=1)

    for order in data:
        codigo_rastreo = order['codigo_rastreo']
        doc.add_heading(f"Número de seguimiento: {codigo_rastreo}", level=2).bold = True

        for product in order['productos']:
            # Obtener información del producto
            talla_info = product['talla']
            talla_parts = talla_info.split(':')
            numero_producto = talla_parts[0]
            color = talla_parts[1].split('(')[0]

            # Construir la ruta de la imagen
            image_folder = os.path.join(base_image_folder, numero_producto, color)
            image_path = os.path.join(image_folder, "1.jpg")

            # Verificar si la imagen existe
            if os.path.exists(image_path):
                # Agregar la imagen y la información del producto
                table = doc.add_table(rows=1, cols=2)
                table.autofit = True
                table.cell(0, 0).width = Inches(2.5)  # Establecer el ancho de la celda para la imagen

                # Agregar la imagen a la celda izquierda
                run = table.cell(0, 0).paragraphs[0].add_run()
                run.add_picture(image_path, width=Inches(2.5))

                # Agregar la información del producto a la celda derecha
                cell_info = table.cell(0, 1).paragraphs[0]
                cell_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                cell_info.add_run(f"Título: {product['titulo']}\n")
                cell_info.add_run(f"Talla: {talla_info}\n")
                cell_info.add_run(f"Cantidad: {product['cantidad']}\n")
                cell_info.add_run(f"Código de Rastreo: {codigo_rastreo}\n")
            else:
                print(f"La imagen {image_path} no existe.")

        # Eliminar espacios adicionales entre los pedidos
        # doc.add_paragraph("--------------------")
        # doc.add_paragraph("\n")  # Eliminado para reducir el espacio entre pedidos

    # Agregar la fecha de generación del documento
    doc.add_paragraph(f"Documento generado el {date.today().strftime('%d/%m/%Y')}")

    # Guardar el archivo Word
    word_file = "pedidos_shein.docx"
    doc.save(word_file)
    print(f"Archivo Word '{word_file}' creado exitosamente.")


def create_tracking_table(data):
    """Crea un archivo Word con la tabla de seguimiento."""
    doc = Document()

    # Agregar título y subtítulo
    doc.add_heading("Hoja de Salida U4U", level=1)
    doc.add_paragraph("Número de seguimiento:", style='Heading2')
    doc.add_paragraph("Fecha de enviado:", style='Heading2')

    # Crear tabla de seguimiento
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    # Establecer estilos para los títulos de las columnas
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Color negro

    # Agregar la información de seguimiento
    for i, order in enumerate(data):
        # Saltar la primera fila ya que contiene los títulos de las columnas
        if i == 0:
            continue

        row = table.add_row().cells
        row[0].text = order['codigo_rastreo']
        fecha_enviado = order.get('fecha_enviado', '')
        row[1].text = fecha_enviado if fecha_enviado else ' '  # Usar espacio en blanco si la fecha es vacía

    # Guardar el archivo Word
    doc_file = "tracking_table.docx"
    doc.save(doc_file)
    print(f"Archivo Word '{doc_file}' creado exitosamente.")


def main():
    # Seleccionar carpeta de imágenes
    image_folder = select_folder()
    if not image_folder:
        print("No se ha seleccionado ninguna carpeta. Saliendo del programa.")
        return

    # Cargar datos del JSON
    with open('datos.json', 'r', encoding='utf-8') as file:
        data = json.load(file)

    # Crear archivo de órdenes
    create_orders_doc(data, image_folder)

    # Crear archivo de tabla de seguimiento
    create_tracking_table(data)

if __name__ == "__main__":
    main()
