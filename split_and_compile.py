import os
import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from docx.shared import Cm, Inches
from docx.enum.table import WD_ROW_HEIGHT_RULE
import tkinter as tk
from tkinter import filedialog
from datetime import datetime
import logging

# Intentar importar docx2pdf para conversión Word->PDF
try:
    from docx2pdf import convert
except ImportError:
    convert = None
    logging.warning("docx2pdf no está instalado. La conversión a PDF se omitirá.")

# Configurar logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# --------------------------- UI ---------------------------------
def ask_brand():
    """
    Selector:
      - Marcas y Licencias  (4 por hoja, sin rotar)
      - Pure and Simple     (4 por hoja, sin rotar)
      - TikTok              (2 por hoja, vertical; sin rotar; la segunda inicia a media hoja)
    """
    sel = {}
    sel_root = tk.Toplevel()
    sel_root.title("Seleccione Marca / Modo de salida")
    sel_root.geometry("320x180")
    sel_var = tk.StringVar(value="Marcas y Licencias")

    tk.Label(sel_root, text="Seleccione MARCA / MODO:").pack(pady=(10, 6))
    tk.Radiobutton(sel_root, text="Marcas y Licencias", variable=sel_var, value="Marcas y Licencias").pack(anchor='w', padx=20)
    tk.Radiobutton(sel_root, text="Pure and Simple",    variable=sel_var, value="Pure and Simple").pack(anchor='w', padx=20)
    tk.Radiobutton(sel_root, text="TikTok (2 por hoja, vertical)", variable=sel_var, value="TikTok").pack(anchor='w', padx=20)

    def on_submit():
        sel['brand'] = sel_var.get()
        sel_root.destroy()

    tk.Button(sel_root, text="OK", command=on_submit).pack(pady=(12, 10))
    sel_root.update_idletasks()
    sel_root.attributes('-topmost', True)
    sel_root.focus_force()
    sel_root.grab_set()
    sel_root.wait_window()
    return sel.get('brand', '')

# --------------------------- LÓGICA ------------------------------
def split_and_compile(root):
    logger.info("▶ Iniciando split_and_compile()")

    # 1) Elige MARCA/MODO
    brand = ask_brand()
    if not brand:
        logger.error("No se seleccionó ninguna marca. Abortando.")
        root.quit()
        return
    logger.info(f"Marca / Modo seleccionado: {brand}")

    # 2) Selecciona PDF
    pdf_path = filedialog.askopenfilename(
        title="Seleccionar archivo PDF",
        filetypes=[("PDF files", "*.pdf")]
    )
    if not pdf_path:
        logger.error("No se seleccionó ningún archivo PDF. Abortando.")
        root.quit()
        return

    # 3) Carpeta de salida
    today = datetime.now().strftime("%d-%m-%Y")
    folder_name = f"Marcas -Shein - {today}"
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    output_folder = os.path.join(desktop, folder_name)
    os.makedirs(output_folder, exist_ok=True)
    logger.info(f"Carpeta de salida: {output_folder}")

    # 4) Abre PDF original
    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        logger.error(f"❌ No se pudo abrir el PDF: {e}")
        root.quit()
        return

    total = doc.page_count
    logger.info(f"Documento original: {total} páginas")

    # --- A) IMPARES → PDF térmico ---
    odd_pdf = fitz.open()
    odd_indices = []
    for i in range(total):
        if (i + 1) % 2 == 1:
            odd_pdf.insert_pdf(doc, from_page=i, to_page=i)
            odd_indices.append(i + 1)

    exp_odds = (total + 1) // 2
    if len(odd_indices) != exp_odds:
        logger.warning(f"Mismatch impares: esperadas={exp_odds}, extraídas={len(odd_indices)} ({odd_indices})")
    else:
        logger.info(f"Páginas impares extraídas: {odd_indices}")

    odd_name = f"{brand} Guias shein {today} - impresora termica.pdf"
    odd_path = os.path.join(output_folder, odd_name)
    odd_pdf.save(odd_path)
    logger.info(f"✅ PDF TÉRMICO guardado en: {odd_path}")

    # --- B) PARES → DOCX según modo ---
    even_indices = [i + 1 for i in range(total) if (i + 1) % 2 == 0]
    logger.info(f"Índices pares a procesar: {even_indices}")

    images = []
    skipped = []
    scale = 2  # factor de zoom para mayor resolución
    for idx in even_indices:
        page = doc.load_page(idx - 1)
        pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        if img.getbbox() is None:
            skipped.append(idx)
            continue

        # En modo TikTok no rotamos (requisito actual)
        images.append((idx, img))

    if skipped:
        logger.warning(f"Páginas pares en blanco omitidas: {skipped}")

    # Crear DOCX tamaño Carta
    doc_word = Document()
    for section in doc_word.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)

    # Tomamos medidas útiles en CM para evitar floats sueltos
    section = doc_word.sections[0]
    usable_w_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    usable_h_cm = section.page_height.cm - section.top_margin.cm - section.bottom_margin.cm

    # Parámetros de maquetación según modo
    if brand == "TikTok":
        per_page = 2
        rows, cols = 2, 1
        picture_width = Cm(usable_w_cm * 0.98)   # casi todo el ancho útil
        out_suffix = "tiktok"
    else:
        per_page = 4
        rows, cols = 2, 2
        picture_width = Cm(7.0)                  # tamaño previo
        out_suffix = "impresora laser"

    count = 0
    table = None
    for idx, img in images:
        if count % per_page == 0:
            table = doc_word.add_table(rows=rows, cols=cols)
            table.autofit = False
            table.allow_autofit = False

            if brand == "TikTok":
                # Forzar que cada fila mida exactamente media página útil
                half_h_cm = usable_h_cm / 2.0
                for r in range(2):
                    row = table.rows[r]
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                    row.height = Cm(half_h_cm)   # <-- ahora es Length, no float

        # Posición dentro de la página
        pos = count % per_page
        row_i = pos // cols
        col_i = pos % cols
        cell = table.cell(row_i, col_i)

        # Insertar imagen a casi todo el ancho (sin rotar)
        temp_png = os.path.join(output_folder, f"even_{idx}.png")
        img.save(temp_png, quality=100)
        run = cell.paragraphs[0].add_run()
        run.add_picture(temp_png, width=picture_width)

        logger.info(f"Par {idx} → hoja {count//per_page + 1}, celda ({row_i},{col_i}), modo {brand}")
        count += 1

        if count % per_page == 0:
            doc_word.add_page_break()

    # Guardar DOCX
    docx_name = f"{brand} Guias shein {today} - {out_suffix}.docx"
    docx_path = os.path.join(output_folder, docx_name)
    doc_word.save(docx_path)
    logger.info(f"✅ DOCX guardado en: {docx_path}")

    # 5) Convertir DOCX → PDF (si disponible)
    if convert:
        pdf_name = f"{brand} Guias shein {today} - {out_suffix}.pdf"
        pdf_path = os.path.join(output_folder, pdf_name)
        convert(docx_path, pdf_path)
        logger.info(f"✅ PDF ({out_suffix}) guardado en: {pdf_path}")
    else:
        logger.error("docx2pdf no disponible; se omitió conversión a PDF de pares.")

    logger.info("▶ Proceso completado exitosamente.")
    root.quit()

def main():
    root = tk.Tk()
    root.iconify()  # mantiene vivo el loop sin mostrar la ventana principal
    root.after(0, split_and_compile, root)
    root.mainloop()

if __name__ == "__main__":
    main()
