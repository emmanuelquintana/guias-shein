from docx import Document
from docx.shared import Cm
from PIL import Image
import os
from datetime import datetime
import tkinter as tk
from tkinter import filedialog
import logging
from docx2pdf import convert

# Configuración de registro (logs)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def select_directory():
    root = tk.Tk()
    root.withdraw()
    directory = filedialog.askdirectory()
    return directory

# Crear un nuevo documento de Word
doc = Document()

# Configurar los márgenes del documento
for section in doc.sections:
    section.left_margin = Cm(0.5)  # Margen izquierdo de 0.5 cm
    section.right_margin = Cm(0.5)  # Margen derecho de 0.5 cm
    section.top_margin = Cm(0.5)  # Margen superior de 0.5 cm
    section.bottom_margin = Cm(0.5)  # Margen inferior de 0.5 cm

# Seleccionar el directorio donde se encuentran las imágenes
logger.info("Selecciona el directorio donde se encuentran las guías.")
directory = select_directory()
logger.info(f"Directorio seleccionado: {directory}")

# Listar los archivos en el directorio
image_files = sorted([f for f in os.listdir(directory) if f.endswith('.jpg')])

# Ancho y alto deseados en centímetros
desired_width = Cm(19.26)
desired_height = Cm(13.22)

# Abrir, redimensionar las imágenes y colocarlas en el documento
logger.info("Agregando las imágenes al documento...")
for i in range(0, len(image_files), 2):
    # Abrir la primera imagen
    img_path1 = os.path.join(directory, image_files[i])
    img1 = Image.open(img_path1)

    # Rotar la imagen 90 grados
    img1 = img1.transpose(Image.ROTATE_270)

    # Calcular el nuevo tamaño manteniendo la relación de aspecto
    img1.thumbnail((desired_height, desired_width))

    # Guardar temporalmente la primera imagen
    img1_temp = "temp_img1.jpg"
    img1.save(img1_temp)

    # Abrir la segunda imagen si existe
    if i + 1 < len(image_files):
        img_path2 = os.path.join(directory, image_files[i + 1])
        img2 = Image.open(img_path2)

        # Rotar la imagen 90 grados
        img2 = img2.transpose(Image.ROTATE_270)

        # Calcular el nuevo tamaño manteniendo la relación de aspecto
        img2.thumbnail((desired_height, desired_width))

        # Guardar temporalmente la segunda imagen
        img2_temp = "temp_img2.jpg"
        img2.save(img2_temp)

        # Agregar las imágenes al documento
        doc.add_picture(img1_temp, width=desired_width, height=desired_height)
        doc.add_picture(img2_temp, width=desired_width, height=desired_height)

        # Eliminar las imágenes temporales
        os.remove(img2_temp)
    else:
        # Agregar solo la primera imagen si no hay segunda imagen
        doc.add_picture(img1_temp, width=desired_width, height=desired_height)

    # Eliminar la imagen temporal
    os.remove(img1_temp)

# Obtener la fecha actual en formato "dd-mm-aaaa"
fecha_actual = datetime.now().strftime("%d-%m-%Y")

# Definir el nombre del archivo con la fecha actual
nombre_archivo_doc = f"Guias Shein {fecha_actual}.docx"
nombre_archivo_pdf = f"Guias Shein {fecha_actual}.pdf"

# Guardar el documento Word con el nombre de archivo especificado
doc.save(nombre_archivo_doc)
logger.info(f"El archivo Word '{nombre_archivo_doc}' ha sido creado con éxito.")

# Convertir el documento Word a PDF
logger.info("Convirtiendo el documento Word a PDF...")
convert(nombre_archivo_doc)
logger.info(f"El archivo PDF '{nombre_archivo_pdf}' ha sido creado con éxito.")
