import fitz
from docx import Document
from docx.shared import Cm
from PIL import Image
import os
import tkinter as tk
from tkinter import filedialog
import logging
from datetime import datetime
from docx2pdf import convert
import math

# Configuración de registro (logs)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def select_pdf_and_convert():
    # Abre una ventana de diálogo para seleccionar el archivo PDF
    root = tk.Tk()
    root.withdraw() # Oculta la ventana principal
    
    file_path = filedialog.askopenfilename(title="Seleccionar archivo PDF", filetypes=[("PDF files", "*.pdf")])
    
    if file_path:
        # Crea el nombre de la carpeta de salida
        output_folder_name = f"GUIAS SHEIN {datetime.now().strftime('%Y-%m-%d')} -IMAGENES"
        
        # Obtiene la ruta completa de la carpeta de salida
        output_folder_path = os.path.join(os.getcwd(), output_folder_name)
        
        # Crea la carpeta de salida si no existe
        if not os.path.exists(output_folder_path):
            os.makedirs(output_folder_path)
        
        # Tamaño mínimo requerido para las imágenes
        min_width = 896
        min_height = 1538
        
        # Convierte el PDF a imágenes JPEG numeradas
        pdf_to_jpg(file_path, output_folder_path, min_width, min_height)
        
        # Ejecuta el proceso de agregar imágenes al documento Word y convertirlo a PDF
        process_images(output_folder_path)

def pdf_to_jpg(pdf_path, output_folder, min_width, min_height):
    # Abre el archivo PDF
    pdf_document = fitz.open(pdf_path)
    
    # Itera sobre cada página del PDF
    for page_number in range(len(pdf_document)):
        # Obtiene la página
        page = pdf_document.load_page(page_number)
        
        # Convierte la página en imagen
        pix = page.get_pixmap()
        
        # Ajusta el tamaño de la imagen si es menor que el mínimo especificado
        if pix.width < min_width or pix.height < min_height:
            scale_factor = max(min_width / pix.width, min_height / pix.height)
            pix = page.get_pixmap(matrix=fitz.Matrix(scale_factor, scale_factor))
        
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        
        # Guarda la imagen como JPEG numerado
        output_path = f"{output_folder}/{page_number + 1:04d}.jpg"  # Formato con ceros iniciales
        img.save(output_path, quality=100)  # Ajusta la calidad al 100%
        logger.info(f"Guardado {output_path}")
    
    # Cierra el archivo PDF
    pdf_document.close()

def process_images(image_folder):
    # Crear un nuevo documento de Word
    doc = Document()

    # Configurar los márgenes del documento
    for section in doc.sections:
        section.left_margin = Cm(0.5)  # Margen izquierdo de 0.5 cm
        section.right_margin = Cm(0.5)  # Margen derecho de 0.5 cm
        section.top_margin = Cm(0.5)  # Margen superior de 0.5 cm
        section.bottom_margin = Cm(0.5)  # Margen inferior de 0.5 cm

    # Seleccionar el directorio donde se encuentran las imágenes
    logger.info("Selecciona el directorio donde se encuentran las guías.")
    directory = image_folder
    logger.info(f"Directorio seleccionado: {directory}")

    # Listar los archivos en el directorio
    image_files = sorted([f for f in os.listdir(directory) if f.endswith('.jpg')])

    # Ancho y alto deseados en centímetros para cada imagen
    desired_width = Cm(7.59)  # Ancho para 2 imágenes en una fila
    desired_height = Cm(13.02)  # Alto para 2 imágenes en una columna

    # Contador para controlar el número de imágenes agregadas por página
    image_count = 0

    # Contador para contar el total de imágenes al inicio
    total_imagenes_inicio = len(image_files)

    # Abrir, redimensionar las imágenes y colocarlas en el documento
    logger.info("Agregando las imágenes al documento...")
    for image_file in image_files:
        # Abrir la imagen
        img_path = os.path.join(directory, image_file)
        img = Image.open(img_path)

        # Calcular el nuevo tamaño manteniendo la relación de aspecto
        img.thumbnail((desired_width, desired_height))

        # Agregar la imagen al documento
        if image_count == 0:
            # Agregar una nueva tabla para las imágenes
            table = doc.add_table(rows=2, cols=2)
            table.autofit = False  # Desactivar ajuste automático del tamaño de la tabla
            table.allow_autofit = False  # Desactivar ajuste automático del tamaño de las celdas

        # Obtener la celda correspondiente para la imagen
        cell = table.cell(image_count // 2, image_count % 2)
        # Agregar la imagen a la celda
        run = cell.paragraphs[0].add_run()
        run.add_picture(img_path, width=desired_width, height=desired_height)

        image_count += 1

        # Si se han agregado 4 imágenes, crear una nueva página y reiniciar el contador de imágenes
        if image_count == 4:
            # Agregar un salto de página
            doc.add_page_break()
            # Reiniciar el contador de imágenes
            image_count = 0

    # Contar el número total de hojas de 4 imágenes
    hojas_de_4_imagenes = math.ceil(total_imagenes_inicio / 4)

    # Obtener la fecha actual en formato "dd-mm-aaaa"
    fecha_actual = datetime.now().strftime("%d-%m-%Y")

    # Definir el nombre del archivo con la fecha actual
    nombre_archivo_doc = f"Guias Shein {fecha_actual} medidas pequeñas4.docx"
    nombre_archivo_pdf = f"Guias Shein {fecha_actual} medidas pequeñas4.pdf"

    # Guardar el documento Word con el nombre de archivo especificado
    doc.save(nombre_archivo_doc)
    logger.info(f"El archivo Word '{nombre_archivo_doc}' ha sido creado con éxito.")

    # Convertir el documento Word a PDF
    logger.info("Convirtiendo el documento Word a PDF...")
    convert(nombre_archivo_doc)
    logger.info(f"El archivo PDF '{nombre_archivo_pdf}' ha sido creado con éxito.")

    # Guardar el PDF generado con las páginas pares eliminadas
    nombre_archivo_pdf_sin_pares = f"Guias Shein {fecha_actual} medidas pequeñas4_sin_pares.pdf"
    eliminar_hojas_pares(nombre_archivo_pdf, nombre_archivo_pdf_sin_pares)
    logger.info(f"El archivo PDF '{nombre_archivo_pdf_sin_pares}' ha sido creado con las páginas pares eliminadas.")

    # Eliminar el PDF original generado por docx2pdf
    os.remove(nombre_archivo_pdf)
    logger.info("El PDF original ha sido eliminado.")

    # Log del total de imágenes al inicio y el número de hojas de 4 imágenes
    logger.info(f"Total de imágenes al inicio: {total_imagenes_inicio}")
    logger.info(f"Total de hojas de 4 imágenes: {hojas_de_4_imagenes}")

def eliminar_hojas_pares(nombre_archivo_pdf, nombre_archivo_pdf_sin_pares):
    # Cargar el documento PDF
    pdf_document = fitz.open(nombre_archivo_pdf)

    # Contador para rastrear el número de páginas eliminadas
    paginas_eliminadas = 0

    # Iterar sobre cada página del PDF en reversa
    for index, _ in reversed(list(enumerate(pdf_document))):
        if (index + 1) % 2 == 0:
            pdf_document.delete_page(index)
            paginas_eliminadas += 1

    # Guardar el documento PDF actualizado
    pdf_document.save(nombre_archivo_pdf_sin_pares)
    pdf_document.close()

    logger.info(f"Se eliminaron {paginas_eliminadas} páginas pares del archivo '{nombre_archivo_pdf_sin_pares}'.")

if __name__ == "__main__":
    select_pdf_and_convert()
