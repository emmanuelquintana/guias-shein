import fitz
from docx import Document
from docx.shared import Cm
from PIL import Image
import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import logging
from datetime import datetime
from docx2pdf import convert
import math
import threading

# Configuración de registro (logs)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def update_progress(root, progress_bar, status_label, progress, message):
    progress_bar['value'] = progress
    status_label.config(text=message)
    root.update_idletasks()

def get_day_from_filename(filename):
    if 'viernes' in filename.lower():
        return 'Viernes'
    elif 'sabado' in filename.lower() or 'sábado' in filename.lower():
        return 'Sábado'
    elif 'domingo' in filename.lower():
        return 'Domingo'
    else:
        return ''

def pdf_to_jpg(pdf_path, output_folder, min_width, min_height, day, pdf_info_text, progress_bar, status_label, root):
    pdf_document = fitz.open(pdf_path)
    total_pages = len(pdf_document)
    pedidos = total_pages // 2
    pdf_info_text.insert(tk.END, f"{os.path.basename(pdf_path)} ({day}): {total_pages} páginas, {pedidos} pedidos\n")

    for page_number in range(total_pages):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        if pix.width < min_width or pix.height < min_height:
            scale_factor = max(min_width / pix.width, min_height / pix.height)
            pix = page.get_pixmap(matrix=fitz.Matrix(scale_factor, scale_factor))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        output_path = f"{output_folder}/{page_number + 1:04d}.jpg"
        img.save(output_path, quality=100)
        logger.info(f"Guardado {output_path}")
        progress = ((page_number + 1) / total_pages) * 100
        update_progress(root, progress_bar, status_label, progress, f"Convirtiendo página {page_number + 1} de {total_pages} del archivo {os.path.basename(pdf_path)}")

    pdf_document.close()

def process_images(image_folder, day, pdf_info_text, progress_bar, status_label, root, prefix=""):
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)

    directory = image_folder
    logger.info(f"Directorio seleccionado: {directory}")
    image_files = sorted([f for f in os.listdir(directory) if f.endswith('.jpg')])
    desired_width = Cm(7.59)
    desired_height = Cm(13.02)
    image_count = 0
    total_imagenes_inicio = len(image_files)
    logger.info("Agregando las imágenes al documento...")
    for image_file in image_files:
        img_path = os.path.join(directory, image_file)
        img = Image.open(img_path)
        img.thumbnail((desired_width, desired_height))
        if image_count == 0:
            table = doc.add_table(rows=2, cols=2)
            table.autofit = False
            table.allow_autofit = False
        cell = table.cell(image_count // 2, image_count % 2)
        run = cell.paragraphs[0].add_run()
        run.add_picture(img_path, width=desired_width, height=desired_height)
        image_count += 1
        if image_count == 4:
            doc.add_page_break()
            image_count = 0

    hojas_de_4_imagenes = math.ceil(total_imagenes_inicio / 4)
    fecha_actual = datetime.now().strftime("%d-%m-%Y")
    nombre_archivo_doc = f"{prefix}Guias Shein {fecha_actual} {day} medidas pequeñas4.docx"
    nombre_archivo_pdf = f"{prefix}Guias Shein {fecha_actual} {day} medidas pequeñas4.pdf"
    doc.save(nombre_archivo_doc)
    logger.info(f"El archivo Word '{nombre_archivo_doc}' ha sido creado con éxito.")
    update_progress(root, progress_bar, status_label, 100, "Convirtiendo el documento Word a PDF...")
    convert(nombre_archivo_doc)
    logger.info(f"El archivo PDF '{nombre_archivo_pdf}' ha sido creado con éxito.")
    nombre_archivo_pdf_sin_pares = f"{prefix}Guias Shein {fecha_actual} {day} medidas pequeñas canguros.pdf"
    eliminar_hojas_pares(nombre_archivo_pdf, nombre_archivo_pdf_sin_pares, progress_bar, status_label, root)
    logger.info(f"El archivo PDF '{nombre_archivo_pdf_sin_pares}' ha sido creado con las páginas pares eliminadas.")
    os.remove(nombre_archivo_pdf)
    logger.info("El PDF original ha sido eliminado.")
    logger.info(f"Total de imágenes al inicio: {total_imagenes_inicio}")
    logger.info(f"Total de hojas de 4 imágenes: {hojas_de_4_imagenes}")
    logger.info(f"Total de Pedidos procesados: {total_imagenes_inicio / 2}")
    process_remaining_images_as_large(image_folder, day, pdf_info_text, progress_bar, status_label, root, prefix)

def eliminar_hojas_pares(nombre_archivo_pdf, nombre_archivo_pdf_sin_pares, progress_bar, status_label, root):
    pdf_document = fitz.open(nombre_archivo_pdf)
    paginas_eliminadas = 0
    for index, _ in reversed(list(enumerate(pdf_document))):
        if (index + 1) % 2 == 0:
            pdf_document.delete_page(index)
            paginas_eliminadas += 1
    pdf_document.save(nombre_archivo_pdf_sin_pares)
    pdf_document.close()
    logger.info(f"Se eliminaron {paginas_eliminadas} páginas pares del archivo '{nombre_archivo_pdf_sin_pares}'.")

def ask_to_process_another(root):
    response = messagebox.askyesno("Proceso completado", "¿Desea procesar otro pedido?")
    if response:
        reset_interface()
    else:
        root.quit()

def reset_interface():
    pdf_info_text.delete(1.0, tk.END)
    progress_bar['value'] = 0
    status_label.config(text="Estado: Esperando selección de archivos")

def select_pdf_and_convert():
    def process_single_day():
        file_path = filedialog.askopenfilename(title="Seleccionar archivo PDF", filetypes=[("PDF files", "*.pdf")])
        if file_path:
            process_files([file_path], is_weekend=False)

    def process_weekend():
        file_paths = filedialog.askopenfilenames(title="Seleccionar archivos PDF (Viernes, Sábado y Domingo)", filetypes=[("PDF files", "*.pdf")])
        if len(file_paths) == 3:
            process_files(file_paths, is_weekend=True, prefix="", folder_label="Marcas")
        else:
            messagebox.showerror("Error", "Debe seleccionar tres archivos PDF (Viernes, Sábado y Domingo).")

    def process_weekend_ps():
        file_paths = filedialog.askopenfilenames(title="Seleccionar archivos PDF (Viernes, Sábado y Domingo)", filetypes=[("PDF files", "*.pdf")])
        if len(file_paths) == 3:
            process_files(file_paths, is_weekend=True, prefix="PS ", folder_label="Pure And Simple")
        else:
            messagebox.showerror("Error", "Debe seleccionar tres archivos PDF (Viernes, Sábado y Domingo).")

    def process_files(file_paths, is_weekend, prefix="", folder_label="Marcas"):
        def run_processing():
            if is_weekend:
                for file_path in file_paths:
                    day = get_day_from_filename(file_path)
                    output_folder_name = f"{folder_label} - GUIAS SHEIN {datetime.now().strftime('%Y-%m-%d')} - {day} - IMAGENES"
                    output_folder_path = os.path.join(os.getcwd(), output_folder_name)
                    if not os.path.exists(output_folder_path):
                        os.makedirs(output_folder_path)
                    min_width = 896
                    min_height = 1538
                    pdf_to_jpg(file_path, output_folder_path, min_width, min_height, day, pdf_info_text, progress_bar, status_label, root)
                    process_images(output_folder_path, day, pdf_info_text, progress_bar, status_label, root, prefix)
                ask_to_process_another(root)
            else:
                file_path = file_paths[0]
                day = get_day_from_filename(file_path)
                output_folder_name = f"{folder_label} - GUIAS SHEIN {datetime.now().strftime('%Y-%m-%d')} - {day} - IMAGENES"
                output_folder_path = os.path.join(os.getcwd(), output_folder_name)
                if not os.path.exists(output_folder_path):
                    os.makedirs(output_folder_path)
                min_width = 896
                min_height = 1538
                pdf_to_jpg(file_path, output_folder_path, min_width, min_height, day, pdf_info_text, progress_bar, status_label, root)
                process_images(output_folder_path, day, pdf_info_text, progress_bar, status_label, root, prefix)
                ask_to_process_another(root)

        threading.Thread(target=run_processing).start()

    global root, pdf_info_text, progress_bar, status_label
    root = tk.Tk()
    root.title("Procesador de Pedidos GUIAS SHEIN")
    root.geometry("700x500")
    root.configure(bg="#f0f4f8")

    style = ttk.Style()
    style.theme_use('clam')
    style.configure('TButton', font=('Segoe UI', 12, 'bold'), foreground='#fff', background='#0078d7', padding=10)
    style.map('TButton', background=[('active', '#005a9e')])
    style.configure('TLabel', font=('Segoe UI', 11), background='#f0f4f8')
    style.configure('TProgressbar', thickness=20)

    title_label = tk.Label(root, text="Procesador de Pedidos GUIAS SHEIN", font=("Segoe UI", 18, "bold"), bg="#f0f4f8", fg="#0078d7")
    title_label.pack(pady=(20, 10))

    single_day_button = ttk.Button(root, text="Seleccionar PDF para un solo día", command=process_single_day)
    single_day_button.pack(pady=10, ipadx=10, ipady=5)

    weekend_button = ttk.Button(root, text="Seleccionar PDFs para el fin de semana (Marcas)", command=process_weekend)
    weekend_button.pack(pady=10, ipadx=10, ipady=5)

    weekend_ps_button = ttk.Button(root, text="Seleccionar PDFs para el fin de semana (Pure And Simple)", command=process_weekend_ps)
    weekend_ps_button.pack(pady=10, ipadx=10, ipady=5)

    progress_bar = ttk.Progressbar(root, orient='horizontal', length=600, mode='determinate', style='TProgressbar')
    progress_bar.pack(pady=25)

    status_label = ttk.Label(root, text="Estado: Esperando selección de archivos", style='TLabel')
    status_label.pack(pady=10)

    pdf_info_text = tk.Text(root, height=10, width=80, font=("Consolas", 10))
    pdf_info_text.pack(pady=10)

    root.mainloop()

def process_remaining_images_as_large(image_folder, day, pdf_info_text, progress_bar, status_label, root, prefix=""):
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)

    directory = image_folder
    logger.info(f"Directorio seleccionado: {directory}")
    image_files = sorted([f for f in os.listdir(directory) if f.endswith('.jpg')])
    desired_width = Cm(19.26)
    desired_height = Cm(13.22)
    logger.info("Agregando las imágenes al documento...")
    for i in range(0, len(image_files), 2):
        img_path1 = os.path.join(directory, image_files[i])
        img1 = Image.open(img_path1)
        img1 = img1.transpose(Image.ROTATE_270)
        img1.thumbnail((desired_height, desired_width))
        img1_temp = "temp_img1.jpg"
        img1.save(img1_temp)
        if i + 1 < len(image_files):
            img_path2 = os.path.join(directory, image_files[i + 1])
            img2 = Image.open(img_path2)
            img2 = img2.transpose(Image.ROTATE_270)
            img2.thumbnail((desired_height, desired_width))
            img2_temp = "temp_img2.jpg"
            img2.save(img2_temp)
            doc.add_picture(img1_temp, width=desired_width, height=desired_height)
            doc.add_picture(img2_temp, width=desired_width, height=desired_height)
            os.remove(img2_temp)
        else:
            doc.add_picture(img1_temp, width=desired_width, height=desired_height)
        os.remove(img1_temp)

    fecha_actual = datetime.now().strftime("%d-%m-%Y")
    nombre_archivo_doc = f"{prefix}Guias Shein {fecha_actual} {day} medidas grandes.docx"
    nombre_archivo_pdf = f"{prefix}Guias Shein {fecha_actual} {day} medidas grandes.pdf"
    doc.save(nombre_archivo_doc)
    logger.info(f"El archivo Word '{nombre_archivo_doc}' ha sido creado con éxito.")
    convert(nombre_archivo_doc)
    logger.info(f"El archivo PDF '{nombre_archivo_pdf}' ha sido creado con éxito.")

if __name__ == "__main__":
    select_pdf_and_convert()
