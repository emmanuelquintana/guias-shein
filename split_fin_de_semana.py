import os
import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from docx.shared import Cm, Inches
import tkinter as tk
from tkinter import filedialog, messagebox
from datetime import datetime
import logging

# Intentar importar docx2pdf para conversión Word->PDF
try:
    from docx2pdf import convert
except ImportError:
    convert = None
    logging.warning("docx2pdf no está instalado. La conversión a PDF se omitirá.")

# Configurar logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

def ask_brand():
    sel = {}
    sel_root = tk.Toplevel()
    sel_root.title("Selecciona la Marca")
    sel_root.geometry("340x180")
    sel_root.configure(bg="#f5f5f5")

    sel_var = tk.StringVar(value="Marcas y Licencias")

    title = tk.Label(sel_root, text="¿Qué marca vas a procesar?", font=("Segoe UI", 13, "bold"), bg="#f5f5f5", fg="#333")
    title.pack(pady=(18, 8))

    frame = tk.Frame(sel_root, bg="#f5f5f5")
    frame.pack(pady=(0, 10))

    rb1 = tk.Radiobutton(frame, text="Marcas y Licencias", variable=sel_var, value="Marcas y Licencias",
                         font=("Segoe UI", 11), bg="#f5f5f5", fg="#222", selectcolor="#e0e0e0", anchor="w", width=20)
    rb1.grid(row=0, column=0, sticky="w", padx=18, pady=2)
    rb2 = tk.Radiobutton(frame, text="Pure and Simple", variable=sel_var, value="Pure and Simple",
                         font=("Segoe UI", 11), bg="#f5f5f5", fg="#222", selectcolor="#e0e0e0", anchor="w", width=20)
    rb2.grid(row=1, column=0, sticky="w", padx=18, pady=2)

    def on_submit():
        sel['brand'] = sel_var.get()
        sel_root.destroy()

    btn = tk.Button(sel_root, text="Continuar", command=on_submit, font=("Segoe UI", 11, "bold"), bg="#4caf50", fg="white", relief="flat", width=14)
    btn.pack(pady=(8, 12))

    sel_root.update_idletasks()
    sel_root.attributes('-topmost', True)
    sel_root.focus_force()
    sel_root.grab_set()
    sel_root.wait_window()

    return sel.get('brand', '')

def ask_weekend():
    root = tk.Toplevel()
    root.title("¿Fin de semana?")
    root.geometry("370x160")
    root.configure(bg="#f5f5f5")
    msg = tk.Label(root, text="¿Es para el fin de semana?\n\nSi eliges Sí, deberás seleccionar 3 PDFs:\nViernes, Sábado y Domingo.",
                   font=("Segoe UI", 11), bg="#f5f5f5", fg="#333", justify="center")
    msg.pack(pady=(18, 10))

    result = {'val': False}
    def yes():
        result['val'] = True
        root.destroy()
    def no():
        result['val'] = False
        root.destroy()

    btn_frame = tk.Frame(root, bg="#f5f5f5")
    btn_frame.pack(pady=(0, 10))
    tk.Button(btn_frame, text="Sí", command=yes, font=("Segoe UI", 11, "bold"), bg="#2196f3", fg="white", width=8, relief="flat").pack(side="left", padx=18)
    tk.Button(btn_frame, text="No", command=no, font=("Segoe UI", 11, "bold"), bg="#e53935", fg="white", width=8, relief="flat").pack(side="left", padx=18)

    root.update_idletasks()
    root.attributes('-topmost', True)
    root.focus_force()
    root.grab_set()
    root.wait_window()
    return result['val']

def ask_pdf_for_day(day_label=None):
    title = f"Seleccionar archivo PDF de {day_label}" if day_label else "Seleccionar archivo PDF"
    return filedialog.askopenfilename(
        title=title,
        filetypes=[("PDF files", "*.pdf")]
    )

def ensure_output_folder(brand):
    today = datetime.now().strftime("%d-%m-%Y")
    safe_brand = brand.replace(" ", "_")
    folder_name = f"{safe_brand} -Shein - {today}"
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    output_folder = os.path.join(desktop, folder_name)
    os.makedirs(output_folder, exist_ok=True)
    return output_folder, today

def process_pdf_for_day(pdf_path, brand, output_folder, today, day_label=None):
    """
    Procesa un PDF con la lógica original. Si day_label está presente,
    lo agrega a los nombres de los archivos de salida.
    """
    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        logger.error(f"❌ No se pudo abrir el PDF ({day_label or 'Único'}): {e}")
        return

    total = doc.page_count
    logger.info(f"[{day_label or 'Único'}] Documento original: {total} páginas")

    # --- A) IMPARES → PDF térmico ---
    odd_pdf = fitz.open()
    odd_indices = []
    for i in range(total):
        if (i + 1) % 2 == 1:
            odd_pdf.insert_pdf(doc, from_page=i, to_page=i)
            odd_indices.append(i + 1)

    exp_odds = (total + 1) // 2
    if len(odd_indices) != exp_odds:
        logger.warning(f"[{day_label or 'Único'}] Mismatch impares: esperadas={exp_odds}, extraídas={len(odd_indices)} ({odd_indices})")
    else:
        logger.info(f"[{day_label or 'Único'}] Páginas impares extraídas: {odd_indices}")

    day_chunk = f" - {day_label}" if day_label else ""

    odd_name = f"{brand} Guias shein {today}{day_chunk} - impresora termica.pdf"
    odd_path = os.path.join(output_folder, odd_name)
    odd_pdf.save(odd_path)
    odd_pdf.close()
    logger.info(f"✅ [{day_label or 'Único'}] PDF TÉRMICO guardado en: {odd_path}")

    # --- B) PARES → DOCX (4 por hoja) + PDF láser ---
    even_indices = [i + 1 for i in range(total) if (i + 1) % 2 == 0]
    logger.info(f"[{day_label or 'Único'}] Índices pares que deberían procesarse: {even_indices}")

    images = []
    skipped = []
    scale = 2  # factor de zoom para mayor resolución
    for idx in even_indices:
        page = doc.load_page(idx - 1)
        pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        if img.getbbox() is None:
            skipped.append(idx)
            continue
        images.append((idx, img))

    if skipped:
        logger.warning(f"[{day_label or 'Único'}] Páginas pares en blanco omitidas: {skipped}")
    if len(images) != len(even_indices) - len(skipped):
        logger.warning(
            f"[{day_label or 'Único'}] Después de filtrar blancas: esperadas={len(even_indices)}, "
            f"procesadas={len(images)}"
        )
    else:
        logger.info(f"[{day_label or 'Único'}] Imágenes pares procesadas correctamente: {[i for i,_ in images]}")

    # Crear DOCX tamaño Carta
    doc_word = Document()
    for section in doc_word.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)

    # Incrustar imágenes 4 por hoja (ligeramente más pequeñas)
    desired_w = Cm(7.0)
    desired_h = Cm(12.0)
    count = 0
    table = None

    for idx, img in images:
        if count % 4 == 0:
            table = doc_word.add_table(rows=2, cols=2)
            table.autofit = False
            table.allow_autofit = False

        row = (count % 4) // 2
        col = (count % 4) % 2
        cell = table.cell(row, col)

        # Temp PNG en la carpeta del día (o principal si no hay día)
        temp_png = os.path.join(
            output_folder,
            f"{(day_label or 'unico').lower()}_even_{idx}.png"
        )
        img.save(temp_png, quality=100)
        cell.paragraphs[0].add_run().add_picture(
            temp_png,
            width=desired_w,
            height=desired_h
        )
        logger.info(f"[{day_label or 'Único'}] Insertada página par {idx} en tabla posición ({row},{col})")
        count += 1

        if count % 4 == 0:
            doc_word.add_page_break()

    # Guardar DOCX
    docx_name = f"{brand} Guias shein {today}{day_chunk} - impresora laser.docx"
    docx_path = os.path.join(output_folder, docx_name)
    doc_word.save(docx_path)
    logger.info(f"✅ [{day_label or 'Único'}] DOCX guardado en: {docx_path}")

    # Convertir DOCX → PDF láser
    if convert:
        try:
            laser_name = f"{brand} Guias shein {today}{day_chunk} - impresora laser.pdf"
            laser_path = os.path.join(output_folder, laser_name)
            convert(docx_path, laser_path)
            logger.info(f"✅ [{day_label or 'Único'}] PDF LÁSER guardado en: {laser_path}")
        except Exception as e:
            logger.error(f"❌ [{day_label or 'Único'}] Error al convertir DOCX a PDF: {e}")
    else:
        logger.error(f"[{day_label or 'Único'}] docx2pdf no disponible; se omitió conversión a PDF láser.")

    doc.close()

def split_and_compile(root):
    logger.info("▶ Iniciando split_and_compile()")

    # 1) Elige MARCA
    brand = ask_brand()
    if not brand:
        logger.error("No se seleccionó ninguna marca. Abortando.")
        root.quit()
        return
    logger.info(f"Marca seleccionada: {brand}")

    # 2) ¿Es para fin de semana?
    weekend = ask_weekend()

    # 3) Prepara carpeta de salida (una sola para todo el proceso)
    output_folder, today = ensure_output_folder(brand)
    logger.info(f"Carpeta de salida: {output_folder}")

    if weekend:
        # Pedimos 3 PDFs: Viernes, Sabado y Domingo
        days = ["Viernes", "Sabado", "Domingo"]
        paths = {}
        for d in days:
            p = ask_pdf_for_day(d)
            if not p:
                logger.error(f"No se seleccionó PDF para {d}. Abortando.")
                root.quit()
                return
            paths[d] = p

        # Crear subcarpetas por día y procesar
        for d in days:
            day_dir = os.path.join(output_folder, d)
            os.makedirs(day_dir, exist_ok=True)
            logger.info(f"▶ Procesando {d} en carpeta: {day_dir}")
            process_pdf_for_day(paths[d], brand, day_dir, today, day_label=d)

        logger.info("▶ Proceso de fin de semana completado exitosamente.")
    else:
        # Flujo normal: un solo PDF en la carpeta principal
        pdf_path = ask_pdf_for_day()
        if not pdf_path:
            logger.error("No se seleccionó ningún archivo PDF. Abortando.")
            root.quit()
            return

        process_pdf_for_day(pdf_path, brand, output_folder, today, day_label=None)
        logger.info("▶ Proceso simple completado exitosamente.")

    root.quit()

def main():
    root = tk.Tk()
    root.iconify()  # mantiene vivo el loop sin mostrar la ventana principal
    root.after(0, split_and_compile, root)
    root.mainloop()

if __name__ == "__main__":
    main()
