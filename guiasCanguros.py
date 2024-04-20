from docx import Document
from docx.shared import Cm
from PIL import Image
import os
from datetime import datetime
import tkinter as tk
from tkinter import filedialog
import logging
from docx2pdf import convert
import PyPDF2
import math

# Configuración de registro (logs)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def select_directory():
    root = tk.Tk()
    root.withdraw()
    directory = filedialog.askdirectory()
    return directory

def remove_even_pages(pdf_path, output_path):
    # Crear un objeto PDFWriter
    pdf_writer = PyPDF2.PdfWriter()

    # Abrir el PDF
    with open(pdf_path, 'rb') as pdf_file:
        # Crear un objeto PDFReader
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Iterar sobre las páginas y agregar solo las páginas impares al PDFWriter
        for page_num in range(len(pdf_reader.pages)):
            if page_num % 2 == 0:
                page = pdf_reader.pages[page_num]
                pdf_writer.add_page(page)

        # Guardar el PDF generado sin las páginas pares
        with open(output_path, 'wb') as output_pdf:
            pdf_writer.write(output_pdf)

# Crear un nuevo documento de Word
doc = Document()

# Configurar los márgenes del documento
for section in doc.sections:
    section.left_margin = Cm(0.5)  # Margen izquierdo de 0.5 cm
    section.right_margin = Cm(0.5)  # Margen derecho de 0.5 cm
    section.top_margin = Cm(0.5)  # Margen superior de 0.5 cm
    section.bottom_margin = Cm(0.5)  # Margen inferior de 0.5 cm

# Seleccionar el directorio donde se encuentran las imágenes
logger.info("Selecciona el directorio donde se encuentran las guías.")
directory = select_directory()
logger.info(f"Directorio seleccionado: {directory}")

# Listar los archivos en el directorio
image_files = sorted([f for f in os.listdir(directory) if f.endswith('.jpg')])

# Contar el número total de imágenes al inicio
total_imagenes_inicio = len(image_files)
logger.info(f"Se encontraron un total de {total_imagenes_inicio} imágenes en el directorio.")
total_hojas = math.ceil(total_imagenes_inicio / 4)
logger.info(f"Total de hojas a generar: {total_hojas}")

# Ancho y alto deseados en centímetros para cada imagen
desired_width = Cm(7.59)  # Ancho para 2 imágenes en una fila
desired_height = Cm(13.02)   # Alto para 2 imágenes en una columna

# Contador para controlar el número de imágenes agregadas por página
image_count = 0

# Contador para las hojas de 4 imágenes
hojas_de_4_imagenes = 0

# Abrir, redimensionar las imágenes y colocarlas en el documento
logger.info("Agregando las imágenes al documento...")
for image_file in image_files:
    # Abrir la imagen
    img_path = os.path.join(directory, image_file)
    img = Image.open(img_path)

    # Calcular el nuevo tamaño manteniendo la relación de aspecto
    img.thumbnail((desired_width, desired_height))

    # Agregar la imagen al documento
    if image_count == 0:
        # Agregar una nueva tabla para las imágenes
        table = doc.add_table(rows=2, cols=2)
        table.autofit = False  # Desactivar ajuste automático del tamaño de la tabla
        table.allow_autofit = False  # Desactivar ajuste automático del tamaño de las celdas

    # Obtener la celda correspondiente para la imagen
    cell = table.cell(image_count // 2, image_count % 2)
    # Agregar la imagen a la celda
    run = cell.paragraphs[0].add_run()
    run.add_picture(img_path, width=desired_width, height=desired_height)

    image_count += 1

    # Si se han agregado 4 imágenes, crear una nueva página y reiniciar el contador de imágenes
    if image_count == 4:
        # Incrementar el contador de hojas de 4 imágenes
        hojas_de_4_imagenes += 1
        # Agregar un salto de página
        doc.add_page_break()
        # Reiniciar el contador de imágenes
        image_count = 0

# Contar el número de imágenes agregadas al final
total_imagenes_agregadas = len(image_files) - total_imagenes_inicio
logger.info(f"Se agregaron un total de {total_imagenes_agregadas} imágenes al documento.")

# Obtener la fecha actual en formato "dd-mm-aaaa"
fecha_actual = datetime.now().strftime("%d-%m-%Y")

# Definir el nombre del archivo con la fecha actual
nombre_archivo_doc = f"Guias Shein {fecha_actual} medidas pequeñas4.docx"
nombre_archivo_pdf = f"Guias Shein {fecha_actual} medidas pequeñas4.pdf"

# Guardar el documento Word con el nombre de archivo especificado
doc.save(nombre_archivo_doc)
logger.info(f"El archivo Word '{nombre_archivo_doc}' ha sido creado con éxito.")

# Convertir el documento Word a PDF
logger.info("Convirtiendo el documento Word a PDF...")
convert(nombre_archivo_doc)
logger.info(f"El archivo PDF '{nombre_archivo_pdf}' ha sido creado con éxito.")

# Guardar el PDF generado con las páginas pares eliminadas
nombre_archivo_pdf_sin_pares = f"Guias Shein {fecha_actual} medidas pequeñas4_sin_pares.pdf"
remove_even_pages(nombre_archivo_pdf, nombre_archivo_pdf_sin_pares)
logger.info(f"El archivo PDF '{nombre_archivo_pdf_sin_pares}' ha sido creado con las páginas pares eliminadas.")

# Eliminar el PDF original generado por docx2pdf
os.remove(nombre_archivo_pdf)
logger.info("El PDF original ha sido eliminado.")

# Log del total de imágenes al inicio y el número de hojas de 4 imágenes
logger.info(f"Total de imágenes al inicio: {total_imagenes_inicio}")
logger.info(f"Total de hojas de 4 imágenes: {hojas_de_4_imagenes}")
